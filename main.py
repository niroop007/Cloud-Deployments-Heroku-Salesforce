# -*- coding: utf-8 -*-
"""
Created on Apr-11-2021

@author: nsriram
"""
import os
from docx import Document
import docx
#from docx2pdf import convert
from flask import Flask, render_template, request, redirect, flash, url_for, send_file
from datetime import date
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE
#import pythoncom
import time

UPLOAD_FOLDER = 'uploads'






today = date.today()
sys_date = today.strftime("%d/%m/%Y")
#print("d1 =", d1)




app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = "secret key"


@app.route('/')
def home():
    
    return render_template('index.html')


#@app.route('/index.html')
#def returnhome():
    
#    return render_template('index.html')

@app.route('/formtodocx',methods = ['GET', 'POST'])
def formtodocx():
    if request.method == 'POST':
        
        client_name=request.form['clientname']
        #
        print(client_name)
        community_name=request.form['commname']
        print(community_name)
        email=request.form['email']
        print(email)
        contact=request.form['contact']
        print(contact)
        address=request.form['address']
        print(address)
        has_b1=request.form['has_b1']
        print(has_b1)
        count_b1=request.form['count_b1']
        print(count_b1)
        #has_b2=request.form['has_b2']
        #count_b2=request.form['count_b2']
        #has_b3=request.form['has_b3']
        #count_b3=request.form['count_b3']
       # watt_low=request.form['9w']
        #watt_high=request.form['12w']
        #o1_price=elsa_price_o1(quantity)
        
        if has_b1=='B1':
            print("Checking first conditions...")
            o1_count=count_b1
            o1_total_price=elsa_price_o1(o1_count)
            o2_total_price=elsa_price_o2(o1_count)
            
        #elif has_b1!='B1' and has_b2=='B2' and has_b3!='B3':
            #o1_count=count_b2
            #o1_total_price=elsa_price_o1(o1_count)
            #o2_total_price=elsa_price_o2(o1_count)
        #elif has_b1!='B1' and has_b2!='B2' and has_b3=='B3':
            #o1_count=count_b3
           # o1_total_price=elsa_price_o1(o1_count)
           # o2_total_price=elsa_price_o2(o1_count)
        #elif has_b1=='B1' and has_b2=='B2' and has_b3!='B3':
        #    o1_count=count_b1+count_b2
         #   o1_total_price=elsa_price_o1(o1_count)
         #   o2_total_price=elsa_price_o2(o1_count)
        #elif has_b1=='B1' and has_b2=='B2' and has_b3=='B3':
        #    o1_count=count_b1+count_b2+count_b3
        #    o1_total_price=elsa_price_o1(o1_count)
        #    o2_total_price=elsa_price_o2(o1_count)
        
        
        #pythoncom.CoInitialize()
        doc = docx.Document(os.path.join(app.config['UPLOAD_FOLDER'],"template.docx"))
        
        
        p=doc.add_paragraph("\n\n\n\n\n\n\n")
        p.add_run('To,\n')
        p.add_run(community_name).bold = True
        #p.add_run("Botanika").bold=True
        p=doc.add_paragraph(address)
        p=doc.add_paragraph(email)
        p=doc.add_paragraph(contact)
        p=doc.add_paragraph("\n\n")
        p=doc.add_paragraph("Dear "+client_name+" Garu,")
        p=doc.add_paragraph("\n")
        p=doc.add_paragraph("The below quote is based on the following details we obtained from you on "+sys_date+".The quote would be adjusted as and when new details/updates are observed during the course of discussion/consultation.")
        p=doc.add_paragraph("\n")
        p=doc.add_paragraph("\t""1. B1 basement is common for everyone and they have 190 plus LED tube lights, which are turned on 24X7.\n"
                            "\t2. Based on our field study we are going install ELSA sensors to around 70% of lights.\n"
                            "\t3. We will leave 30% of tube lights to maintain…\n"
                            "\t\t\ta. Ambient lighting\n"
                            "\t\t\tb. Driveways\n"
                            "\t4. The number of sensors need to be installed may vary based on the ground truth. The invoice would be adjusted, to reflect the final number of sensors installed.\n"
                            "\t5. As discussed and suggested, installation would be done by UDAK Team.\n"
                            "\t6. Payment of One Time Model would be in 2 installments (50% on delivery and 50% on completion of installation).\n"
                            "\t7. Payment of Subscription Model would be post-paid and will start from the date of delivery.\n")
        
        
        doc.add_page_break()
        p=doc.add_paragraph("\n\n\n\n")
        #p.add_run('Quote\n').bold = True
        p=doc.add_paragraph("The following equipment and charges would be part of the final quotation submitted.\n")
        
        
        ################ONE TIME PAYMENT########################
        p=doc.add_paragraph("OPTION 1: One Time Payment\n")
        table1 = doc.add_table(rows = 1, cols = 4, style='Table Grid')
        OTP_cells = table1.rows[0].cells
        OTP_cells[0].text = 'Item'
        OTP_cells[1].text = 'Description'
        OTP_cells[2].text = 'Quantity'
        OTP_cells[3].text = 'Price (Rs.)'
        cells = table1.add_row().cells
        cells[0].text = "1."
        cells[1].text = "ELSA sensor with inbuilt activity and lux detection, along with adjustable knobs for specific adjustments. (Rs. 799/Sensor)*"
        cells[2].text = o1_count
        cells[3].text = o1_total_price
        cells = table1.add_row().cells
        cells[0].text = "2."
        cells[1].text = "One Year Full Product Replacement Guarantee**"
        cells[2].text = " "
        cells[3].text = "Free"

        ##########################################################
        
        
       

        ################Subscription Model########################
        
        p=doc.add_paragraph("\nOPTION 2: Subscription Model\n")
        table2 = doc.add_table(rows = 1, cols = 4, style='Table Grid')
        SM_cells = table2.rows[0].cells
        SM_cells[0].text = 'Item'
        SM_cells[1].text = 'Description'
        SM_cells[2].text = 'Quantity'
        SM_cells[3].text = 'Price (Rs.)'
        cells = table2.add_row().cells
        cells[0].text = "1."
        cells[1].text = "Solution is subscription based, wherein company gives the product free and users pay nominal monthly fee of Rs 75/Sensor*"
        cells[2].text = o1_count
        cells[3].text = o2_total_price
        cells = table2.add_row().cells
        cells[0].text = "2."
        cells[1].text = "Full Product Maintenance & Replacement Guarantee**"
        cells[2].text = " "
        cells[3].text = "Lifetime Free***"
        ###########################################################
        p=doc.add_paragraph("\n\n* 18% GST is exclusive of the above details and would be part of the final invoice.\n"
                             "** Any and every defective sensor would be fully replaced, with no additional charges to the customer.\n"
                             "*** Product will be owned & maintained by UDAK under active subscription\n"
                             "\nIf you have any questions, please do not hesitate to contact us.\n"
                             "\n\nKind regards,\n"
                             "UDAK Technologies Pvt Ltd.")
        
        
        
        
        
        
        
        
        
        
        doc.save(os.path.join("\\static\\" +client_name+ ".pdf"))
        time.sleep(10)
        doc_path=".\\static\\"+client_name+".pdf"
        print(doc_path)
        print(os.path.getsize(doc_path))
	#print(os.path.getsize(doc_path))
        #docxtopdf(doc_path)
        quote_path=client_name+".pdf"
        print(quote_path)
        #file_download(quote_path)
    return render_template('result.html', filename=quote_path)
    #flash(client_name)
    
        
       



def generate_smatrix(docx_file):
    convert(docx_file)


#def docxtopdf(docx_file):
#    convert(docx_file)
#    time.sleep(5)
#    print("PDF Created Successfully")
    
def elsa_price_o1(total1):
    return total1*799

def elsa_price_o2(total2):
    return total2*79
    

#@app.route('/file_download', methods = ['GET', 'POST'])
#def file_download(filename):
#    print("Entered File_download method & Sending file path to HTML")
#    return send_file(filename)
    
    #return render_template('result.html')
if __name__ == '__main__':
	app.run()
